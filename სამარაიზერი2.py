#!/usr/bin/env python
# coding: utf-8

# In[45]:


pip install python-docx


# In[2]:


#pip install pdfplumber


# In[1]:


import pdfplumber


# In[2]:


with pdfplumber.open("/Users/tatiatsiklauri/Desktop/mrm-151-04eng-2.pdf") as pdf:
    text = ""
    for page in pdf.pages:
        text += page.extract_text()


# In[3]:


from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer


# In[6]:


parser = PlaintextParser.from_string(text, Tokenizer("english"))
summarizer = LsaSummarizer()
summary = summarizer(parser.document, sentences_count=50)


# In[7]:


#print("Original Text:")
#print(text)
print("\nSummary:")
for sentence in summary:
    print(sentence)


# In[37]:


import docx


# In[38]:


doc = docx.Document()


# In[39]:


###doc.add_heading('Original Text:', 0)
###doc.add_paragraph(text)


# In[40]:


doc.add_heading('Summary:', 0)


# In[41]:


for sentence in summary: doc.add_paragraph(str(sentence))


# In[42]:


doc.save('/Users/tatiatsiklauri/Desktop/summary.docx')


# In[ ]:





# In[ ]:




